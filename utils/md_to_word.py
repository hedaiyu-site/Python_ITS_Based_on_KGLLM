from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# 读取Markdown文件
def read_markdown(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

# 解析Markdown并转换为Word文档
def convert_md_to_word(markdown_content, output_path):
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 解析Markdown内容
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 处理标题
        if line.startswith('#'):
            level = line.count('#')
            title_text = line[level:].strip()
            
            if level == 1:
                # 一级标题
                para = doc.add_heading(title_text, level=0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                # 二级标题
                doc.add_heading(title_text, level=1)
            elif level == 3:
                # 三级标题
                doc.add_heading(title_text, level=2)
        
        # 处理表格
        elif line.startswith('|'):
            # 查找表格的所有行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # 回退一步，因为循环结束后会i++
            
            # 解析表格
            if len(table_lines) >= 3:  # 至少需要表头、分隔线和一行数据
                # 创建表格
                header_row = table_lines[0].split('|')[1:-1]  # 移除首尾的|，分割成列表
                header_cells = [cell.strip() for cell in header_row]
                
                # 计算表格列数
                col_count = len(header_cells)
                
                # 创建表格，行数为数据行数+1（表头）
                data_rows = table_lines[2:]
                table = doc.add_table(rows=1 + len(data_rows), cols=col_count)
                
                # 设置表头
                header_cells_word = table.rows[0].cells
                for j, cell_text in enumerate(header_cells):
                    header_cells_word[j].text = cell_text
                
                # 填充数据
                for row_idx, data_line in enumerate(data_rows):
                    data_cells = data_line.split('|')[1:-1]
                    data_cells = [cell.strip() for cell in data_cells]
                    
                    # 确保数据单元格数量与表头一致
                    if len(data_cells) < col_count:
                        data_cells.extend([''] * (col_count - len(data_cells)))
                    elif len(data_cells) > col_count:
                        data_cells = data_cells[:col_count]
                    
                    # 填充数据行
                    word_cells = table.rows[row_idx + 1].cells
                    for j, cell_text in enumerate(data_cells):
                        # 处理链接
                        link_match = re.search(r'\[(.*?)\]\((.*?)\)', cell_text)
                        if link_match:
                            # 暂时只保留链接文本，不处理链接功能
                            word_cells[j].text = link_match.group(1)
                        else:
                            word_cells[j].text = cell_text
        
        # 处理列表
        elif line.startswith('- '):
            # 查找连续的列表项
            list_items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                list_items.append(lines[i].strip()[2:])  # 移除'- '前缀
                i += 1
            i -= 1  # 回退一步
            
            # 添加项目符号列表
            for item in list_items:
                para = doc.add_paragraph(item, style='List Bullet')
        
        # 处理普通段落
        elif line:
            # 处理特殊标记，如✅
            line = line.replace('✅', '✓')
            doc.add_paragraph(line)
        
        # 处理空行
        else:
            doc.add_paragraph()
        
        i += 1
    
    # 保存文档
    doc.save(output_path)
    print(f"Word文档已保存到: {output_path}")

# 主函数
if __name__ == "__main__":
    input_file = "文档\知识图谱设计与构建报告.md"
    output_file = "文档\知识图谱设计与构建报告.docx"
    
    markdown_content = read_markdown(input_file)
    convert_md_to_word(markdown_content, output_file)
